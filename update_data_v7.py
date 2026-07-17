"""
Riscrive nei documenti .docx salvati in Documenti/ (contratti e revisioni) il titolo/riferimento
progetto in modo che corrisponda esattamente al nome progetto ora assegnato in app al contratto
collegato. Un file puo' essere condiviso da piu' contratti/progetti (riutilizzato come demo per piu'
call-off dello stesso cliente): in tal caso si usa il progetto del contratto con id piu' basso
(quello "originario") e lo si segnala in output.
Eseguire una sola volta: py update_data_v7.py
"""
import sys
from collections import defaultdict

sys.path.insert(0, r'C:\Users\bs_udeschini\Desktop\Progetto\Avanzamento-contratti')
from app import app, Allegato, Variante
from docx import Document

def imposta_testo_paragrafo(paragrafo, nuovo_testo):
    if not paragrafo.runs:
        paragrafo.add_run(nuovo_testo)
        return
    paragrafo.runs[0].text = nuovo_testo
    for r in paragrafo.runs[1:]:
        r.text = ''


with app.app_context():
    mappa_file = defaultdict(list)  # percorso -> [(contratto_id, progetto_nome)]

    for a in Allegato.query.filter(Allegato.percorso.like('%.docx')).all():
        if a.contratto.progetto:
            mappa_file[a.percorso].append((a.contratto.id, a.contratto.progetto.nome))

    for v in Variante.query.all():
        if v.percorso_allegato and v.contratto.progetto:
            mappa_file[v.percorso_allegato].append((v.contratto.id, v.contratto.progetto.nome))

    modificati = 0
    for percorso, associazioni in mappa_file.items():
        associazioni_uniche = sorted(set(associazioni))
        progetto_canonico = associazioni_uniche[0][1]

        doc = Document(percorso)
        if not doc.paragraphs:
            continue

        primo_paragrafo = doc.paragraphs[0]
        testo_originale = primo_paragrafo.text.strip()

        if testo_originale.upper().startswith('ATTO DI MODIFICA'):
            # Documento di revisione: aggiunge il progetto come sottotitolo se non già presente
            secondo = doc.paragraphs[1] if len(doc.paragraphs) > 1 else None
            marcatore = f'Progetto: {progetto_canonico}'
            if secondo and marcatore not in secondo.text:
                nuovo_paragrafo = secondo.insert_paragraph_before(marcatore)
                if secondo.runs:
                    nuovo_paragrafo.runs[0].font.bold = True if nuovo_paragrafo.runs else None
        else:
            imposta_testo_paragrafo(primo_paragrafo, progetto_canonico.upper())

        doc.save(percorso)
        modificati += 1

        nome_file = percorso.split('\\')[-1]
        extra = f'  (file condiviso anche da: {[c for c, _ in associazioni_uniche[1:]]})' \
            if len(associazioni_uniche) > 1 else ''
        print(f'{nome_file:45s} -> {progetto_canonico}{extra}')

    print(f'\nDocumenti aggiornati: {modificati}')
