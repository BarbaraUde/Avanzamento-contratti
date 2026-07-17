import io
import os
import mimetypes
from datetime import date, datetime
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, abort
from flask_sqlalchemy import SQLAlchemy
import pandas as pd
from docx import Document as DocxDocument

app = Flask(__name__)
app.config['SECRET_KEY'] = 'metra-contratti-2024'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///contratti.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024

db = SQLAlchemy(app)

# Cartella di rete/condivisa in cui sono depositati i documenti (contratti, offerte,
# ordini, revisioni...). Usata dal selettore file per attingere ai documenti esistenti
# senza doverli ricaricare manualmente.
DOCS_ROOT = r'C:\Users\bs_udeschini\Desktop\Progetto\Documenti'
DOCS_ESTENSIONI = {'.pdf', '.docx', '.doc', '.xlsx', '.xls',
                   '.png', '.jpg', '.jpeg', '.gif', '.webp'}


# =============================================================================
# MODELLI
# =============================================================================

class Cliente(db.Model):
    __tablename__ = 'clienti'
    id = db.Column(db.Integer, primary_key=True)
    nome = db.Column(db.String(200), nullable=False)
    codice = db.Column(db.String(50))
    note = db.Column(db.Text)
    progetti = db.relationship('Progetto', backref='cliente', lazy='select',
                               cascade='all, delete-orphan')
    contratti = db.relationship('Contratto', backref='cliente', lazy='select',
                                cascade='all, delete-orphan')
    righe_fatturato = db.relationship('RigaFatturato',
                                      foreign_keys='RigaFatturato.cliente_id',
                                      backref='cliente', lazy='select')


class Progetto(db.Model):
    __tablename__ = 'progetti'
    id = db.Column(db.Integer, primary_key=True)
    cliente_id = db.Column(db.Integer, db.ForeignKey('clienti.id'), nullable=False)
    nome = db.Column(db.String(200), nullable=False)
    codice = db.Column(db.String(50))
    descrizione = db.Column(db.Text)
    stato = db.Column(db.String(50), default='aperto')
    contratti = db.relationship('Contratto', backref='progetto', lazy='select')
    righe_fatturato = db.relationship('RigaFatturato',
                                      foreign_keys='RigaFatturato.progetto_id',
                                      backref='progetto', lazy='select')


class Contratto(db.Model):
    __tablename__ = 'contratti'
    id = db.Column(db.Integer, primary_key=True)
    cliente_id = db.Column(db.Integer, db.ForeignKey('clienti.id'), nullable=False)
    progetto_id = db.Column(db.Integer, db.ForeignKey('progetti.id'), nullable=True)
    codice = db.Column(db.String(100))
    descrizione = db.Column(db.Text)
    valore_totale = db.Column(db.Float, default=0.0)
    stato = db.Column(db.String(50), default='aperto')
    data_apertura = db.Column(db.Date)
    data_chiusura = db.Column(db.Date)
    note = db.Column(db.Text)
    ordini = db.relationship('Ordine', backref='contratto', lazy='select',
                             cascade='all, delete-orphan')
    varianti = db.relationship('Variante', backref='contratto', lazy='select',
                               cascade='all, delete-orphan')
    allegati = db.relationship('Allegato', backref='contratto', lazy='select',
                               cascade='all, delete-orphan')
    righe_fatturato = db.relationship('RigaFatturato',
                                      foreign_keys='RigaFatturato.contratto_id',
                                      backref='contratto', lazy='select')
    articoli = db.relationship('Articolo', backref='contratto', lazy='select',
                               cascade='all, delete-orphan')

    @property
    def num_varianti(self):
        return len(self.varianti)

    @property
    def valore_spedito(self):
        return sum(r.importo or 0.0 for r in self.righe_fatturato)

    @property
    def percentuale_avanzamento(self):
        if self.valore_totale and self.valore_totale > 0:
            return min(100.0, round(self.valore_spedito / self.valore_totale * 100, 1))
        return 0.0

    @property
    def valore_residuo(self):
        return max(0.0, (self.valore_totale or 0.0) - self.valore_spedito)


class Ordine(db.Model):
    __tablename__ = 'ordini'
    id = db.Column(db.Integer, primary_key=True)
    contratto_id = db.Column(db.Integer, db.ForeignKey('contratti.id'), nullable=False)
    numero_ordine = db.Column(db.String(100))
    data = db.Column(db.Date)
    importo = db.Column(db.Float, default=0.0)
    stato = db.Column(db.String(50), default='aperto')
    note = db.Column(db.Text)
    percorso_allegato = db.Column(db.String(1000))


class Variante(db.Model):
    __tablename__ = 'varianti'
    id = db.Column(db.Integer, primary_key=True)
    contratto_id = db.Column(db.Integer, db.ForeignKey('contratti.id'), nullable=False)
    numero = db.Column(db.Integer, nullable=False)
    data = db.Column(db.Date)
    descrizione = db.Column(db.Text)
    valore_delta = db.Column(db.Float, default=0.0)
    percorso_allegato = db.Column(db.String(1000))


class Allegato(db.Model):
    __tablename__ = 'allegati'
    id = db.Column(db.Integer, primary_key=True)
    contratto_id = db.Column(db.Integer, db.ForeignKey('contratti.id'), nullable=False)
    tipo = db.Column(db.String(50), default='contratto')
    nome = db.Column(db.String(300), nullable=False)
    percorso = db.Column(db.String(1000), nullable=False)
    data_aggiunta = db.Column(db.DateTime, default=datetime.utcnow)
    note = db.Column(db.Text)


class Articolo(db.Model):
    __tablename__ = 'articoli'
    id = db.Column(db.Integer, primary_key=True)
    contratto_id = db.Column(db.Integer, db.ForeignKey('contratti.id'), nullable=False)
    codice = db.Column(db.String(100))
    descrizione = db.Column(db.Text)
    quantita_totale = db.Column(db.Float, default=0.0)
    quantita_spedita = db.Column(db.Float, default=0.0)
    unita_misura = db.Column(db.String(20), default='pz')
    prezzo_unitario = db.Column(db.Float, default=0.0)
    note = db.Column(db.Text)
    percorso_allegato = db.Column(db.String(1000))

    @property
    def stato_spedizione(self):
        if self.quantita_totale and self.quantita_spedita >= self.quantita_totale:
            return 'completo'
        if self.quantita_spedita and self.quantita_spedita > 0:
            return 'parziale'
        return 'da spedire'

    @property
    def percentuale_spedita(self):
        if self.quantita_totale and self.quantita_totale > 0:
            return min(100.0, round(self.quantita_spedita / self.quantita_totale * 100, 1))
        return 0.0

    @property
    def valore_articolo(self):
        return (self.quantita_totale or 0.0) * (self.prezzo_unitario or 0.0)

    @property
    def valore_spedito(self):
        return (self.quantita_spedita or 0.0) * (self.prezzo_unitario or 0.0)


class RigaFatturato(db.Model):
    __tablename__ = 'righe_fatturato'
    id = db.Column(db.Integer, primary_key=True)
    contratto_id = db.Column(db.Integer, db.ForeignKey('contratti.id'), nullable=True)
    cliente_id = db.Column(db.Integer, db.ForeignKey('clienti.id'), nullable=True)
    progetto_id = db.Column(db.Integer, db.ForeignKey('progetti.id'), nullable=True)
    data_documento = db.Column(db.Date)
    numero_documento = db.Column(db.String(100))
    importo = db.Column(db.Float, default=0.0)
    descrizione = db.Column(db.Text)
    fonte = db.Column(db.String(200))
    percorso_allegato = db.Column(db.String(1000))
    importato_il = db.Column(db.DateTime, default=datetime.utcnow)


# =============================================================================
# FILTRI E CONTEXT PROCESSOR
# =============================================================================

@app.template_filter('euro')
def euro_filter(value):
    if value is None:
        return '—'
    formatted = f'{value:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.')
    return f'€ {formatted}'


@app.template_filter('data_it')
def data_it_filter(value):
    if value is None:
        return '—'
    if isinstance(value, (date, datetime)):
        return value.strftime('%d/%m/%Y')
    return str(value)


@app.template_filter('avanz_class')
def avanz_class_filter(pct):
    if pct >= 80:
        return 'success'
    if pct >= 40:
        return 'warning'
    return 'danger'


@app.context_processor
def inject_globals():
    return {'today': datetime.now().strftime('%d/%m/%Y')}


# =============================================================================
# HELPER
# =============================================================================

def parse_date(s):
    if not s:
        return None
    for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y'):
        try:
            return datetime.strptime(s.strip(), fmt).date()
        except (ValueError, AttributeError):
            continue
    return None


def parse_float(s):
    if not s:
        return 0.0
    s = str(s).replace(' ', '').replace('.', '').replace(',', '.')
    try:
        return float(s)
    except ValueError:
        return 0.0


# =============================================================================
# ROUTES — DASHBOARD
# =============================================================================

@app.route('/')
def dashboard():
    clienti = Cliente.query.order_by(Cliente.nome).all()
    cliente_id = request.args.get('cliente_id', type=int)
    progetto_id = request.args.get('progetto_id', type=int)
    stato_f = request.args.get('stato', '')

    progetti = Progetto.query.filter_by(cliente_id=cliente_id).order_by(Progetto.nome).all() \
        if cliente_id else Progetto.query.order_by(Progetto.nome).all()

    q = Contratto.query
    if cliente_id:
        q = q.filter_by(cliente_id=cliente_id)
    if progetto_id:
        q = q.filter_by(progetto_id=progetto_id)
    if stato_f:
        q = q.filter_by(stato=stato_f)
    contratti = q.order_by(Contratto.id.desc()).all()

    tot_valore = sum(c.valore_totale or 0 for c in contratti)
    tot_spedito = sum(c.valore_spedito for c in contratti)
    tot_residuo = sum(c.valore_residuo for c in contratti)
    avg_avanz = round(tot_spedito / tot_valore * 100, 1) if tot_valore > 0 else 0.0

    return render_template('dashboard.html',
                           contratti=contratti, clienti=clienti, progetti=progetti,
                           cliente_id=cliente_id, progetto_id=progetto_id, stato_f=stato_f,
                           tot_valore=tot_valore, tot_spedito=tot_spedito,
                           tot_residuo=tot_residuo, avg_avanz=avg_avanz)


# =============================================================================
# ROUTES — CLIENTI
# =============================================================================

@app.route('/clienti')
def clienti_lista():
    clienti = Cliente.query.order_by(Cliente.nome).all()
    return render_template('clienti_lista.html', clienti=clienti)


@app.route('/clienti/nuovo', methods=['GET', 'POST'])
def cliente_nuovo():
    if request.method == 'POST':
        nome = request.form.get('nome', '').strip()
        if not nome:
            flash('Il nome cliente è obbligatorio.', 'danger')
            return redirect(url_for('cliente_nuovo'))
        db.session.add(Cliente(
            nome=nome,
            codice=request.form.get('codice', '').strip(),
            note=request.form.get('note', '').strip()
        ))
        db.session.commit()
        flash(f'Cliente "{nome}" creato.', 'success')
        return redirect(url_for('clienti_lista'))
    return render_template('cliente_form.html', obj=None, title='Nuovo Cliente')


@app.route('/clienti/<int:id>/modifica', methods=['GET', 'POST'])
def cliente_modifica(id):
    c = Cliente.query.get_or_404(id)
    if request.method == 'POST':
        nome = request.form.get('nome', '').strip()
        if not nome:
            flash('Il nome cliente è obbligatorio.', 'danger')
            return redirect(url_for('cliente_modifica', id=id))
        c.nome = nome
        c.codice = request.form.get('codice', '').strip()
        c.note = request.form.get('note', '').strip()
        db.session.commit()
        flash(f'Cliente "{c.nome}" aggiornato.', 'success')
        return redirect(url_for('clienti_lista'))
    return render_template('cliente_form.html', obj=c, title='Modifica Cliente')


@app.route('/clienti/<int:id>/elimina', methods=['POST'])
def cliente_elimina(id):
    c = Cliente.query.get_or_404(id)
    nome = c.nome
    db.session.delete(c)
    db.session.commit()
    flash(f'Cliente "{nome}" eliminato.', 'warning')
    return redirect(url_for('clienti_lista'))


# =============================================================================
# ROUTES — PROGETTI
# =============================================================================

@app.route('/progetti')
def progetti_lista():
    cliente_id = request.args.get('cliente_id', type=int)
    q = Progetto.query
    if cliente_id:
        q = q.filter_by(cliente_id=cliente_id)
    progetti = q.order_by(Progetto.nome).all()
    clienti = Cliente.query.order_by(Cliente.nome).all()
    return render_template('progetti_lista.html', progetti=progetti, clienti=clienti,
                           cliente_id_sel=cliente_id)


@app.route('/progetti/nuovo', methods=['GET', 'POST'])
def progetto_nuovo():
    clienti = Cliente.query.order_by(Cliente.nome).all()
    if request.method == 'POST':
        nome = request.form.get('nome', '').strip()
        cliente_id = request.form.get('cliente_id', type=int)
        if not nome or not cliente_id:
            flash('Nome e Cliente sono obbligatori.', 'danger')
            return render_template('progetto_form.html', obj=None, clienti=clienti, title='Nuovo Progetto')
        db.session.add(Progetto(
            nome=nome, cliente_id=cliente_id,
            codice=request.form.get('codice', '').strip(),
            descrizione=request.form.get('descrizione', '').strip(),
            stato=request.form.get('stato', 'aperto')
        ))
        db.session.commit()
        flash(f'Progetto "{nome}" creato.', 'success')
        return redirect(url_for('progetti_lista'))
    return render_template('progetto_form.html', obj=None, clienti=clienti, title='Nuovo Progetto')


@app.route('/progetti/<int:id>/modifica', methods=['GET', 'POST'])
def progetto_modifica(id):
    p = Progetto.query.get_or_404(id)
    clienti = Cliente.query.order_by(Cliente.nome).all()
    if request.method == 'POST':
        nome = request.form.get('nome', '').strip()
        cliente_id = request.form.get('cliente_id', type=int)
        if not nome or not cliente_id:
            flash('Nome e Cliente sono obbligatori.', 'danger')
            return render_template('progetto_form.html', obj=p, clienti=clienti, title='Modifica Progetto')
        p.nome = nome
        p.cliente_id = cliente_id
        p.codice = request.form.get('codice', '').strip()
        p.descrizione = request.form.get('descrizione', '').strip()
        p.stato = request.form.get('stato', 'aperto')
        db.session.commit()
        flash(f'Progetto "{p.nome}" aggiornato.', 'success')
        return redirect(url_for('progetti_lista'))
    return render_template('progetto_form.html', obj=p, clienti=clienti, title='Modifica Progetto')


@app.route('/progetti/<int:id>/elimina', methods=['POST'])
def progetto_elimina(id):
    p = Progetto.query.get_or_404(id)
    nome = p.nome
    db.session.delete(p)
    db.session.commit()
    flash(f'Progetto "{nome}" eliminato.', 'warning')
    return redirect(url_for('progetti_lista'))


# =============================================================================
# ROUTES — CONTRATTI
# =============================================================================

@app.route('/contratti')
def contratti_lista():
    clienti = Cliente.query.order_by(Cliente.nome).all()
    cliente_id = request.args.get('cliente_id', type=int)
    progetto_id = request.args.get('progetto_id', type=int)
    stato_f = request.args.get('stato', '')

    progetti = Progetto.query.filter_by(cliente_id=cliente_id).order_by(Progetto.nome).all() \
        if cliente_id else Progetto.query.order_by(Progetto.nome).all()

    q = Contratto.query
    if cliente_id:
        q = q.filter_by(cliente_id=cliente_id)
    if progetto_id:
        q = q.filter_by(progetto_id=progetto_id)
    if stato_f:
        q = q.filter_by(stato=stato_f)
    contratti = q.order_by(Contratto.id.desc()).all()

    return render_template('contratti_lista.html', contratti=contratti,
                           clienti=clienti, progetti=progetti,
                           cliente_id=cliente_id, progetto_id=progetto_id, stato_f=stato_f)


@app.route('/contratti/nuovo', methods=['GET', 'POST'])
def contratto_nuovo():
    clienti = Cliente.query.order_by(Cliente.nome).all()
    progetti = Progetto.query.order_by(Progetto.nome).all()
    if request.method == 'POST':
        cliente_id = request.form.get('cliente_id', type=int)
        if not cliente_id:
            flash('Il cliente è obbligatorio.', 'danger')
            return render_template('contratto_form.html', obj=None, clienti=clienti,
                                   progetti=progetti, title='Nuovo Contratto')
        c = Contratto(
            cliente_id=cliente_id,
            progetto_id=request.form.get('progetto_id', type=int) or None,
            codice=request.form.get('codice', '').strip(),
            descrizione=request.form.get('descrizione', '').strip(),
            valore_totale=parse_float(request.form.get('valore_totale', '0')),
            stato=request.form.get('stato', 'aperto'),
            data_apertura=parse_date(request.form.get('data_apertura')),
            data_chiusura=parse_date(request.form.get('data_chiusura')),
            note=request.form.get('note', '').strip()
        )
        db.session.add(c)
        db.session.commit()
        flash(f'Contratto "{c.codice or c.id}" creato.', 'success')
        return redirect(url_for('contratto_dettaglio', id=c.id))
    return render_template('contratto_form.html', obj=None, clienti=clienti,
                           progetti=progetti, title='Nuovo Contratto')


@app.route('/contratti/<int:id>')
def contratto_dettaglio(id):
    c = Contratto.query.get_or_404(id)
    varianti_ord = sorted(c.varianti, key=lambda v: v.numero)
    ordini_ord = sorted(c.ordini, key=lambda o: o.data or date.min, reverse=True)
    articoli_ord = sorted(c.articoli, key=lambda a: (a.codice or ''))
    return render_template('contratto_dettaglio.html', c=c,
                           varianti=varianti_ord, ordini=ordini_ord,
                           articoli=articoli_ord)


@app.route('/contratti/<int:id>/modifica', methods=['GET', 'POST'])
def contratto_modifica(id):
    c = Contratto.query.get_or_404(id)
    clienti = Cliente.query.order_by(Cliente.nome).all()
    progetti = Progetto.query.order_by(Progetto.nome).all()
    if request.method == 'POST':
        cliente_id = request.form.get('cliente_id', type=int)
        if not cliente_id:
            flash('Il cliente è obbligatorio.', 'danger')
            return render_template('contratto_form.html', obj=c, clienti=clienti,
                                   progetti=progetti, title='Modifica Contratto')
        c.cliente_id = cliente_id
        c.progetto_id = request.form.get('progetto_id', type=int) or None
        c.codice = request.form.get('codice', '').strip()
        c.descrizione = request.form.get('descrizione', '').strip()
        c.valore_totale = parse_float(request.form.get('valore_totale', '0'))
        c.stato = request.form.get('stato', 'aperto')
        c.data_apertura = parse_date(request.form.get('data_apertura'))
        c.data_chiusura = parse_date(request.form.get('data_chiusura'))
        c.note = request.form.get('note', '').strip()
        db.session.commit()
        flash('Contratto aggiornato.', 'success')
        return redirect(url_for('contratto_dettaglio', id=c.id))
    return render_template('contratto_form.html', obj=c, clienti=clienti,
                           progetti=progetti, title='Modifica Contratto')


@app.route('/contratti/<int:id>/elimina', methods=['POST'])
def contratto_elimina(id):
    c = Contratto.query.get_or_404(id)
    codice = c.codice or str(c.id)
    db.session.delete(c)
    db.session.commit()
    flash(f'Contratto "{codice}" eliminato.', 'warning')
    return redirect(url_for('contratti_lista'))


# =============================================================================
# ROUTES — ORDINI
# =============================================================================

@app.route('/contratti/<int:contratto_id>/ordini/nuovo', methods=['POST'])
def ordine_nuovo(contratto_id):
    Contratto.query.get_or_404(contratto_id)
    db.session.add(Ordine(
        contratto_id=contratto_id,
        numero_ordine=request.form.get('numero_ordine', '').strip(),
        data=parse_date(request.form.get('data')),
        importo=parse_float(request.form.get('importo', '0')),
        stato=request.form.get('stato', 'aperto'),
        note=request.form.get('note', '').strip(),
        percorso_allegato=request.form.get('percorso_allegato', '').strip() or None
    ))
    db.session.commit()
    flash('Ordine aggiunto.', 'success')
    return redirect(url_for('contratto_dettaglio', id=contratto_id) + '#ordini')


@app.route('/ordini/<int:id>/elimina', methods=['POST'])
def ordine_elimina(id):
    o = Ordine.query.get_or_404(id)
    cid = o.contratto_id
    db.session.delete(o)
    db.session.commit()
    flash('Ordine eliminato.', 'warning')
    return redirect(url_for('contratto_dettaglio', id=cid) + '#ordini')


# =============================================================================
# ROUTES — VARIANTI
# =============================================================================

@app.route('/contratti/<int:contratto_id>/varianti/nuovo', methods=['POST'])
def variante_nuova(contratto_id):
    c = Contratto.query.get_or_404(contratto_id)
    next_num = max((v.numero for v in c.varianti), default=0) + 1
    delta = parse_float(request.form.get('valore_delta', '0'))
    db.session.add(Variante(
        contratto_id=contratto_id,
        numero=next_num,
        data=parse_date(request.form.get('data')),
        descrizione=request.form.get('descrizione', '').strip(),
        valore_delta=delta,
        percorso_allegato=request.form.get('percorso_allegato', '').strip()
    ))
    if delta:
        c.valore_totale = (c.valore_totale or 0.0) + delta
        if c.stato == 'aperto':
            c.stato = 'in corso'
    db.session.commit()
    flash(f'Variante {next_num} aggiunta.', 'success')
    return redirect(url_for('contratto_dettaglio', id=contratto_id) + '#varianti')


@app.route('/varianti/<int:id>/elimina', methods=['POST'])
def variante_elimina(id):
    v = Variante.query.get_or_404(id)
    cid = v.contratto_id
    c = v.contratto
    if v.valore_delta:
        c.valore_totale = max(0.0, (c.valore_totale or 0.0) - v.valore_delta)
    db.session.delete(v)
    db.session.commit()
    flash('Variante eliminata.', 'warning')
    return redirect(url_for('contratto_dettaglio', id=cid) + '#varianti')


# =============================================================================
# ROUTES — ALLEGATI
# =============================================================================

@app.route('/contratti/<int:contratto_id>/allegati/nuovo', methods=['POST'])
def allegato_nuovo(contratto_id):
    Contratto.query.get_or_404(contratto_id)
    nome = request.form.get('nome', '').strip()
    percorso = request.form.get('percorso', '').strip()
    if not nome or not percorso:
        flash('Nome e percorso sono obbligatori.', 'danger')
        return redirect(url_for('contratto_dettaglio', id=contratto_id) + '#allegati')
    db.session.add(Allegato(
        contratto_id=contratto_id,
        tipo=request.form.get('tipo', 'contratto'),
        nome=nome,
        percorso=percorso,
        note=request.form.get('note', '').strip()
    ))
    db.session.commit()
    flash('Allegato aggiunto.', 'success')
    return redirect(url_for('contratto_dettaglio', id=contratto_id) + '#allegati')


@app.route('/allegati/<int:id>/elimina', methods=['POST'])
def allegato_elimina(id):
    a = Allegato.query.get_or_404(id)
    cid = a.contratto_id
    db.session.delete(a)
    db.session.commit()
    flash('Allegato eliminato.', 'warning')
    return redirect(url_for('contratto_dettaglio', id=cid) + '#allegati')


# =============================================================================
# ROUTES — ARTICOLI
# =============================================================================

@app.route('/contratti/<int:contratto_id>/articoli/nuovo', methods=['POST'])
def articolo_nuovo(contratto_id):
    Contratto.query.get_or_404(contratto_id)
    db.session.add(Articolo(
        contratto_id=contratto_id,
        codice=request.form.get('codice', '').strip(),
        descrizione=request.form.get('descrizione', '').strip(),
        quantita_totale=parse_float(request.form.get('quantita_totale', '0')),
        quantita_spedita=parse_float(request.form.get('quantita_spedita', '0')),
        unita_misura=request.form.get('unita_misura', 'pz').strip(),
        prezzo_unitario=parse_float(request.form.get('prezzo_unitario', '0')),
        note=request.form.get('note', '').strip(),
        percorso_allegato=request.form.get('percorso_allegato', '').strip() or None
    ))
    db.session.commit()
    flash('Articolo aggiunto.', 'success')
    return redirect(url_for('contratto_dettaglio', id=contratto_id) + '#articoli')


@app.route('/articoli/<int:id>/elimina', methods=['POST'])
def articolo_elimina(id):
    a = Articolo.query.get_or_404(id)
    cid = a.contratto_id
    db.session.delete(a)
    db.session.commit()
    flash('Articolo eliminato.', 'warning')
    return redirect(url_for('contratto_dettaglio', id=cid) + '#articoli')


@app.route('/articoli/<int:id>/aggiorna', methods=['POST'])
def articolo_aggiorna(id):
    a = Articolo.query.get_or_404(id)
    a.quantita_spedita = parse_float(request.form.get('quantita_spedita', str(a.quantita_spedita)))
    db.session.commit()
    flash('Quantità spedita aggiornata.', 'success')
    return redirect(url_for('contratto_dettaglio', id=a.contratto_id) + '#articoli')


# =============================================================================
# ROUTES — FATTURATO (inserimento manuale riga)
# =============================================================================

@app.route('/contratti/<int:contratto_id>/fatturato/nuovo', methods=['POST'])
def riga_fatturato_nuova(contratto_id):
    c = Contratto.query.get_or_404(contratto_id)
    db.session.add(RigaFatturato(
        contratto_id=contratto_id,
        cliente_id=c.cliente_id,
        progetto_id=c.progetto_id,
        data_documento=parse_date(request.form.get('data_documento')),
        numero_documento=request.form.get('numero_documento', '').strip(),
        importo=parse_float(request.form.get('importo', '0')),
        descrizione=request.form.get('descrizione', '').strip(),
        fonte=request.form.get('fonte', '').strip() or 'Inserimento manuale',
        percorso_allegato=request.form.get('percorso_allegato', '').strip() or None
    ))
    db.session.commit()
    flash('Riga di fatturato aggiunta.', 'success')
    return redirect(url_for('contratto_dettaglio', id=contratto_id) + '#fatturato')


@app.route('/fatturato/<int:id>/elimina', methods=['POST'])
def riga_fatturato_elimina(id):
    r = RigaFatturato.query.get_or_404(id)
    cid = r.contratto_id
    db.session.delete(r)
    db.session.commit()
    flash('Riga di fatturato eliminata.', 'warning')
    return redirect(url_for('contratto_dettaglio', id=cid) + '#fatturato') if cid \
        else redirect(url_for('dashboard'))


# =============================================================================
# ROUTE — SELETTORE FILE (attinge alla cartella documenti condivisa)
# =============================================================================

@app.route('/file/browse')
def file_browse():
    """Restituisce l'albero dei documenti disponibili in DOCS_ROOT, raggruppati per cartella."""
    if not os.path.isdir(DOCS_ROOT):
        return jsonify({'cartelle': []})

    cartelle = []
    for cartella_corrente, dirs, files in os.walk(DOCS_ROOT):
        dirs.sort()
        nome_relativo = os.path.relpath(cartella_corrente, DOCS_ROOT)
        nome_visualizzato = 'Documenti' if nome_relativo == '.' else nome_relativo.replace('\\', ' / ')

        elenco_file = []
        for f in sorted(files):
            ext = os.path.splitext(f)[1].lower()
            if ext not in DOCS_ESTENSIONI:
                continue
            percorso_completo = os.path.join(cartella_corrente, f)
            elenco_file.append({
                'nome': f,
                'percorso': percorso_completo,
                'estensione': ext.lstrip('.')
            })

        if elenco_file:
            cartelle.append({'nome': nome_visualizzato, 'file': elenco_file})

    return jsonify({'cartelle': cartelle})


@app.route('/file/serve')
def file_serve():
    path = request.args.get('path', '').strip()
    if not path:
        abort(400)
    path = os.path.normpath(path)
    if not os.path.isfile(path):
        abort(404)
    mime, _ = mimetypes.guess_type(path)
    mime = mime or 'application/octet-stream'
    inline = mime in ('application/pdf', 'image/png', 'image/jpeg',
                      'image/gif', 'image/webp', 'image/svg+xml')
    return send_file(path, mimetype=mime,
                     as_attachment=not inline,
                     download_name=os.path.basename(path))


@app.route('/file/anteprima')
def file_anteprima():
    """Anteprima universale: PDF/immagini in iframe diretto, DOCX/XLSX renderizzati come HTML."""
    path = request.args.get('path', '').strip()
    if not path:
        abort(400)
    path = os.path.normpath(path)
    if not os.path.isfile(path):
        return render_template('file_anteprima.html', errore='File non trovato sul percorso indicato.',
                               nome=os.path.basename(path))

    ext = path.rsplit('.', 1)[-1].lower() if '.' in path else ''
    nome = os.path.basename(path)

    if ext == 'pdf' or ext in ('png', 'jpg', 'jpeg', 'gif', 'webp', 'svg'):
        return redirect(url_for('file_serve', path=path))

    if ext == 'docx':
        try:
            doc = DocxDocument(path)
            blocchi = []
            for para in doc.paragraphs:
                if para.text.strip():
                    stile = (para.style.name or '').lower() if para.style is not None else ''
                    livello = 'heading' if 'heading' in stile or 'title' in stile else 'paragrafo'
                    blocchi.append({'tipo': livello, 'testo': para.text})
            tabelle = []
            for tbl in doc.tables:
                righe = [[cella.text for cella in riga.cells] for riga in tbl.rows]
                tabelle.append(righe)
            return render_template('file_anteprima.html', nome=nome, tipo='docx',
                                   blocchi=blocchi, tabelle=tabelle)
        except Exception as e:
            return render_template('file_anteprima.html', nome=nome,
                                   errore=f'Impossibile leggere il documento: {e}')

    if ext in ('xlsx', 'xls'):
        try:
            fogli = pd.read_excel(path, sheet_name=None)
            fogli_html = {
                nome_foglio: df.fillna('').to_html(classes='table table-sm table-bordered table-hover mb-0',
                                                   index=False, border=0)
                for nome_foglio, df in fogli.items()
            }
            return render_template('file_anteprima.html', nome=nome, tipo='xlsx', fogli=fogli_html)
        except Exception as e:
            return render_template('file_anteprima.html', nome=nome,
                                   errore=f'Impossibile leggere il file Excel: {e}')

    return render_template('file_anteprima.html', nome=nome,
                           errore='Anteprima non disponibile per questo tipo di file. Usa "Apri" per scaricarlo.')


# =============================================================================
# ROUTES — IMPORT EXCEL FATTURATO
# =============================================================================

@app.route('/import', methods=['GET', 'POST'])
def import_excel():
    clienti = Cliente.query.order_by(Cliente.nome).all()
    contratti = Contratto.query.order_by(Contratto.codice).all()

    if request.method == 'POST':
        # Step 2: conferma import con mapping colonne
        if 'col_importo' in request.form:
            df_json = request.form.get('df_json', '')
            try:
                df = pd.read_json(io.StringIO(df_json), orient='split')
            except Exception as e:
                flash(f'Errore parsing dati: {e}', 'danger')
                return redirect(url_for('import_excel'))

            col_data = request.form.get('col_data') or None
            col_importo = request.form.get('col_importo')
            col_documento = request.form.get('col_documento') or None
            col_descrizione = request.form.get('col_descrizione') or None
            col_cliente = request.form.get('col_cliente') or None
            col_contratto = request.form.get('col_contratto') or None
            fonte = request.form.get('filename', 'import Excel')

            clienti_map = {c.nome.lower().strip(): c for c in clienti}
            contratti_map = {(c.codice or '').lower().strip(): c for c in contratti if c.codice}

            importate = non_assoc = 0
            for _, row in df.iterrows():
                try:
                    importo_val = float(
                        str(row.get(col_importo, 0) or 0).replace(',', '.').replace(' ', '')
                    )
                except (ValueError, TypeError):
                    importo_val = 0.0

                cliente_obj = None
                if col_cliente:
                    nome_cl = str(row.get(col_cliente, '') or '').strip().lower()
                    cliente_obj = clienti_map.get(nome_cl)

                contratto_obj = None
                if col_contratto:
                    cod = str(row.get(col_contratto, '') or '').strip().lower()
                    contratto_obj = contratti_map.get(cod)

                data_val = None
                if col_data:
                    raw_d = row.get(col_data)
                    if pd.notna(raw_d):
                        try:
                            if isinstance(raw_d, (date, datetime)):
                                data_val = raw_d.date() if isinstance(raw_d, datetime) else raw_d
                            else:
                                data_val = pd.to_datetime(str(raw_d)).date()
                        except Exception:
                            pass

                db.session.add(RigaFatturato(
                    contratto_id=contratto_obj.id if contratto_obj else None,
                    cliente_id=(cliente_obj.id if cliente_obj
                                else (contratto_obj.cliente_id if contratto_obj else None)),
                    data_documento=data_val,
                    numero_documento=str(row.get(col_documento, '') or '').strip() if col_documento else None,
                    importo=importo_val,
                    descrizione=str(row.get(col_descrizione, '') or '').strip() if col_descrizione else None,
                    fonte=fonte
                ))
                importate += 1
                if not contratto_obj:
                    non_assoc += 1

            db.session.commit()
            flash(f'Importate {importate} righe ({non_assoc} non associate a contratto).', 'success')
            return redirect(url_for('dashboard'))

        # Step 1: lettura file e anteprima colonne
        f = request.files.get('file')
        if not f or not f.filename:
            flash('Nessun file selezionato.', 'danger')
            return redirect(url_for('import_excel'))
        try:
            raw = f.read()
            df = pd.read_excel(io.BytesIO(raw))
        except Exception as e:
            flash(f'Errore lettura file: {e}', 'danger')
            return redirect(url_for('import_excel'))

        colonne = list(df.columns)
        anteprima = df.head(5).fillna('').astype(str).to_dict(orient='records')
        df_json = df.to_json(orient='split', date_format='iso')

        return render_template('import_excel.html', step='mapping',
                               colonne=colonne, anteprima=anteprima,
                               df_json=df_json, filename=f.filename,
                               clienti=clienti, contratti=contratti)

    return render_template('import_excel.html', step='upload',
                           clienti=clienti, contratti=contratti)


# =============================================================================
# API
# =============================================================================

@app.route('/api/progetti')
def api_progetti():
    cliente_id = request.args.get('cliente_id', type=int)
    q = Progetto.query
    if cliente_id:
        q = q.filter_by(cliente_id=cliente_id)
    return jsonify([{'id': p.id, 'nome': p.nome} for p in q.order_by(Progetto.nome).all()])


# =============================================================================
# INIT
# =============================================================================

def _migra_colonne_mancanti():
    """Aggiunge in modo additivo le colonne introdotte dopo la creazione iniziale del DB."""
    from sqlalchemy import text
    colonne_richieste = {
        'ordini': [('percorso_allegato', 'VARCHAR(1000)')],
        'articoli': [('percorso_allegato', 'VARCHAR(1000)')],
        'righe_fatturato': [('percorso_allegato', 'VARCHAR(1000)')],
    }
    with db.engine.connect() as conn:
        for tabella, colonne in colonne_richieste.items():
            esistenti = {row[1] for row in conn.execute(text(f'PRAGMA table_info({tabella})'))}
            for nome_colonna, tipo_sql in colonne:
                if nome_colonna not in esistenti:
                    conn.execute(text(f'ALTER TABLE {tabella} ADD COLUMN {nome_colonna} {tipo_sql}'))
                    conn.commit()


with app.app_context():
    db.create_all()
    _migra_colonne_mancanti()

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
